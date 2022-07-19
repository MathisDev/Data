#!/usr/bin/python3
import os
import cv2
from termios import VSWTCH
from tkinter import *
import tkinter as tk
from tkinter import messagebox
from functools import partial
import sqlite3
import functools
from tkinter import font
from turtle import color, width
from ttkwidgets.autocomplete import AutocompleteCombobox
from tkinter import RIGHT, TOP, ttk
from tkinter.messagebox import askokcancel, showinfo, WARNING
from PIL import ImageTk,Image
from ftplib import FTP
from tkinterdnd2 import *


# Classe de 'stockage' :
class Data():
	
	def __init__(self) :
		self.utility="Stockage"


# - Variable de stockage de chemins -
chemin=""


# Classe de mise en forme :
class main():


	def __init__(self):
		self.fontFil = open("sett.txt","r+")
		size = self.fontFil.read()
		self.fontFil.close()


		# - Script de création de la BDD -
		def creer_base():
			sql =["CREATE TABLE Clients (idClient INTEGER PRIMARY KEY NOT NULL UNIQUE, firmeNom TEXT, firmeAdresse TEXT, firmeCP TEXT, firmeVille TEXT, clientNom TEXT, clientPrenom TEXT, adresseMail TEXT, telephone TEXT, soldeCompte REAL, numSiret INTEGER, numTVA INTEGER);","CREATE TABLE Commandes (idCommande INTEGER PRIMARY KEY NOT NULL UNIQUE, dateDepot TEXT, adresseLivraison TEXT, cp INTEGER, ville TEXT, prixConvenu REAL,devis REAL,acompte REAL,resterPaye REAL, description TEXT, travaf TEXT, photo TEXT, idClient INTEGER NOT NULL, FOREIGN KEY (idClient) REFERENCES Clients(idClient));"]
			li = sqlite3.connect("base.db")
			
			if li:
				cons = li.cursor()

				for i in range(len(sql)):
					cons.execute(sql[i])
				li.commit()
				li.close()


		# - Permet de créer la BDD si elle n'exite pas déjà -
		def verifier_base(base):
			from os.path import isfile,isdir
			from os import makedirs

			if not isfile(base) and not isdir('factures'):
				creer_base()
				makedirs('factures')
				
			else:
				print("Base existe déjà !")



		# - Appel de la fonction verifier_base à l'initialisation de l'UI -		
		verifier_base('base.db')

		def toggleFullScreen(event):
			self.fullScreenState = not self.fullScreenState
			window.attributes("-fullscreen", self.fullScreenState)

		def quitFullScreen(event):
			self.fullScreenState = False
			window.attributes("-fullscreen", self.fullScreenState)
			window.geometry('1500x1000')
		
		
		barr_rech = Data()
		barr_rechC = Data()
		barr_rechCli = Data()
	
		window = TkinterDnD.Tk()
		window.attributes('-fullscreen', True)
		fullScreenState = False
		window.bind('<Escape>', quitFullScreen)
		window.bind('<Control-p>', toggleFullScreen)
		window.title("Logiciel Gestion Commande")
		window.geometry('1500x1000')
		window.configure(bg="#666363")

		#self.titrerestor = tk.Label(window,text=" REST'OR ",font =('normal',size,'bold'),bg="#666363").place(x=700,y=15)
		self.titrerestor = tk.Label(window,text=" REST'OR ",font =('normal',20,'bold'),bg="#666363").place(x=700,y=15)

		
		# - Affiche tout les clients -
		def showAllClients():

			color1 = "#8E44AD"
			color = 'grey'
			fg = 13	
			
			self.topShowAllClients = tk.Toplevel()
			self.topShowAllClients.geometry("1920x1080")
			self.topShowAllClients.bind("<Escape>", lambda even: self.topShowAllClients.destroy())

			arbre = ttk.Treeview(self.topShowAllClients, column=("c1","c2", "c3","c4","c5","c6","c7","c8","c9"), show='headings')
			
			conn = sqlite3.connect('base.db')
			cur = conn.cursor()
			rqt = "SELECT firmeNom, firmeAdresse, firmeCP, firmeVille, clientNom, clientPrenom, adresseMail, telephone, soldeCompte FROM Clients;"
			cur.execute(rqt)
			rows = cur.fetchall()
			for row in rows:
				print(row) 
				arbre.insert("", tk.END, values=row)
				arbre.bind("<Double-1>",showClients)
			cur.close()
			conn.close()
			
			l=["firmeNom","firmeAdresse","firmeCP","firmeVille","clientNom","clientPrenom","adresseMail","telephone","soldeCompte"]

			for i in range(len(l)):
				arbre.column("#"+str(i+1), anchor=tk.CENTER)
				arbre.heading("#"+str(i+1), text=l[i])
			arbre.pack(side = "bottom", expand=True, fill='y')

			barr_rechCli.queryString = AutocompleteCombobox(self.topShowAllClients,width=30,font=('Times', 18),completevalues=recuperer("client"))
			barr_rechCli.queryString.pack(padx=10,pady=130)
			barr_rechCli.queryString.bind("<Return>", lambda even, brci = barr_rechCli, wC = self.topShowAllClients: search(brci,wC,"client"))
			
			self.topShowAllClients.bind("<Return>", lambda even, brci = barr_rechCli, wC = self.topShowAllClients: search(brci,wC,"client"))
			tk.Button(self.topShowAllClients, text='Chercher',highlightbackground = "#666363", command=partial(search,barr_rechCli,self.topShowAllClients,"client")).place(x=900,y=130)
			tk.Button(self.topShowAllClients, text='Quitter', command=lambda: self.topShowAllClients.destroy()).place(x=650,y=200)
			

		# - Permet d'effacer un client ET UNE COMMANDE !-
		def unset_item(event):
			tree = event.widget
			citem = tree.focus()
			val = tree.item(citem)
			print(len(val['values']))
			if len(val['values']) == 9:
				answer = askokcancel(
				title='confirmation',
				message='êtes vous sûr de supprimer le client ?',
				icon=WARNING
				)
				if answer:
					nom = val["values"][4]
					prenom = val["values"][5]
					print(nom)
					print(prenom)
					conn = sqlite3.connect('base.db')
					cur = conn.cursor()
					cur.execute('DELETE FROM Commandes WHERE idClient = (SELECT idClient FROM Clients WHERE clientNom = ? and clientPrenom = ?);', (nom, prenom) )
					conn.commit()
					cur.execute('DELETE FROM Clients WHERE clientNom = ?;', [nom])
					conn.commit()
					conn.close()
					showinfo(
					title='Suppression',
					message='Le client à été supprimé'
					)
			else:
				answer = askokcancel(
				title='confirmation',
				message='êtes vous sûr de supprimer la commande ?',
				icon=WARNING
				)
				if answer:
					idComm = val['values'][0]
					conn = sqlite3.connect('base.db')
					cur = conn.cursor()
					cur.execute('DELETE FROM Commandes WHERE idCommande = ?', [idComm])
					conn.commit()
					conn.close()
					showinfo(
					title='Suppression',
					message='La commande à été supprimée'
					)
					
		# - Recherche un client par son nom de famille ou une commande par son numéro -
		def search(entree, window, CI):
			from datetime import datetime
			import re
			conn = sqlite3.connect("base.db")
			
			try :
				if datetime.strptime(entree.queryString.get(), "%d-%m-%Y") is not None:
					print("date ok")
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7"), show='headings')
					l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description"]
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					if entree.queryString != None:
						print(entree.queryString)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient WHERE Commandes.dateDepot = ?;", [entree.queryString.get()])
					fetch = cur.fetchall()
					print("fonte" + str(fetch))
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-3>", unset_item)
						tree.bind("<Button-3>", unset_item)						
					cur.close()
					return 1
				
			except:
				pass

			try:
				reg = re.compile(".*\..*")
				if reg.match(str(entree.queryString.get())):
					print("real")
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7"), show='headings')
					l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description"]
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					if entree.queryString != None:
						print(entree.queryString)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient WHERE Commandes.prixConvenu = ?;", [entree.queryString.get()])
					fetch = cur.fetchall()
					print("fonte" + str(fetch))
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-3>", unset_item)
						tree.bind("<Button-3>", unset_item)						
					cur.close()
					
					return 1 
			except:
				pass
				
			try:
				reg = re.compile(".* .*")
				if reg.match(str(entree.queryString.get())):
					print("desc")
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7"), show='headings')
					l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description"]
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					if entree.queryString != None:
						print(entree.queryString)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient WHERE Commandes.description = ? or Commandes.travaf = ?;", (entree.queryString.get(),entree.queryString.get()))
					fetch = cur.fetchall()
					print("fonte" + str(fetch))
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-3>", unset_item)
						tree.bind("<Button-3>", unset_item)						
					cur.close()
					
					
					return 1 
			except:
				pass
					

			try:

				if type(int(entree.queryString.get()))==int and CI == "general" or CI == "commande":
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7"), show='headings')
					l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description"]
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					if entree.queryString != None:
						print(entree.queryString)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient WHERE idCommande = ?;", [entree.queryString.get()])
					fetch = cur.fetchall()
					print("fonte" + str(fetch))
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-3>", unset_item)
						tree.bind("<Button-3>", unset_item)						
					cur.close()

			except ValueError:
				if CI == "general" or CI == "client":
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7","c8","c9"), show='headings')
					l = ["firmeNom","firmeAdresse","firmeCP","firmeVille","clientNom","clientPrenom","adresseMail","telephone","soldeCompte"]			
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					if entree.queryString != None:
						print(entree.queryString)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT firmeNom, firmeAdresse, firmeCP, firmeVille, clientNom, clientPrenom, adresseMail, telephone, soldeCompte FROM Clients WHERE clientNom LIKE ?;", [entree.queryString.get()])
					fetch = cur.fetchall()
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-3>", unset_item)	
						tree.bind("<Button-2>", unset_item)	
					cur.close()
			
				conn.close()

		def nouvCommande(event):
			tree = event.widget
			citem = tree.focus()
			val = tree.item(citem)
			#print('Valeurs: ' + val + "\nLongueur: " + len(val))
			#print(len(val['values']))
			if len(val['values']) == 9:
				nom = val["values"][4]
				prenom = val["values"][5]

				print("Nom:" + str(nom) + "\nPrenom:" + str(prenom))
			
				conn = sqlite3.connect('base.db')
				cur = conn.cursor()
				cur.execute('SELECT idClient FROM Clients WHERE clientNom = ? and clientPrenom = ?;', (nom,prenom) )
				conn.commit()
				idCli = cur.fetchall()

				cur.execute('SELECT adresseLivraison,cp,ville FROM Commandes WHERE idClient = ?;', [idCli[0][0]])
				Comm = cur.fetchall()
				print(Comm)
			
				windowCommande(self, idCli, Comm)
			
			elif len(val['values']) == 8 or len(val['values']) == 7:
				print(val['values'][0])
				
				conn = sqlite3.connect('base.db')
				cur = conn.cursor()
				cur.execute('SELECT * FROM Commandes WHERE idCommande = ?;', [val['values'][0]])
				conn.commit()
				fetch = cur.fetchall()
				val = fetch[0]
				idCLI = fetch[0][11]
				windowCommande(self, idCLI, val)
				
			
		# - Récupération des noms des clients -
		def recuperer(CI):
			print(CI)
			conn = sqlite3.connect("base.db")
			cur = conn.cursor()
			cur.execute('SELECT clientNom FROM Clients')
			fetch = cur.fetchall()
			cur.close()
			l = [r[0] for r in fetch]

			cur = conn.cursor()
			cur.execute('SELECT idCommande FROM Commandes')
			fetch1 = cur.fetchall()
			cur.close()
			conn.close()
			l1 = [str(r[0]) for r in fetch1]

			if CI == "general":
				return l+l1
			elif CI == "client":
				return l
			elif CI == "commande":
				return l1

		# - Barre de recherche et tableau -
		barr_rech.queryString = AutocompleteCombobox(window,width=30,font=('Times', 18),completevalues=recuperer("general"))
		barr_rech.queryString.pack(padx=10,pady=130)
		tk.Button(window, text='Chercher',highlightbackground = "#666363", command=partial(search,barr_rech,window,"general")).place(x=900,y=130)
		
		barr_rech.queryString.bind("<Return>", lambda even, brci = barr_rech, w = window: search(brci,w,"general"))
		window.bind("<Return>", lambda even, brci = barr_rech, w = window: search(brci,w,"general"))
			
		
		
		# - Affiche tout les clients -
		def showAllCommandes():
			
			self.topShowAllCommandes = tk.Toplevel()
			self.topShowAllCommandes.geometry("1920x1080")
			self.topShowAllCommandes.bind("<Escape>", lambda even: self.topShowAllCommandes.destroy())

			arbre = ttk.Treeview(self.topShowAllCommandes, column=("c1", "c2","c3","c4","c5","c6","c7","c8"), show='headings')
			conn = sqlite3.connect('base.db')
			cur = conn.cursor()
			cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description,travaf FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient ;")
			rows = cur.fetchall()
			for row in rows:
				print(row) 
				arbre.insert("", tk.END, values=row)
				arbre.bind("<Double-1>", nouvCommande)
				arbre.bind("<Button-3>", unset_item)
				arbre.bind("<Button-2>", unset_item)
			cur.close()
			conn.close()
			

			l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description","Travail à faire"]

			for i in range(len(l)):
				arbre.column("#"+str(i+1), anchor=tk.CENTER)
				arbre.heading("#"+str(i+1), text=l[i])
			arbre.pack(side="bottom",expand=True, fill='y')		
			
			
			barr_rechC.queryString = AutocompleteCombobox(self.topShowAllCommandes,width=30,font=('Times', 18),completevalues=recuperer("commande"))
			barr_rechC.queryString.pack(padx=10,pady=130)
			
			barr_rechC.queryString.bind("<Return>", lambda even, brci = barr_rechC, wC = self.topShowAllCommandes: search(brci,wC,"commande"))
			self.topShowAllCommandes.bind("<Return>", lambda even, brci = barr_rechC, wC = self.topShowAllCommandes: search(brci,wC,"commande"))
			
			tk.Button(self.topShowAllCommandes, text='Chercher',highlightbackground = "#666363", command=partial(search,barr_rechC,self.topShowAllCommandes,"commande")).place(x=900,y=130)
			tk.Button(self.topShowAllCommandes, text='Quitter', command=lambda: self.topShowAllCommandes.destroy()).place(x=650,y=200)


		
		# - Permet d'ajouter une commande -
		def windowCommande(self, idC, Comm):

			global chemin
			print(chemin)
			
			from datetime import datetime

			l = 25
			color = "#52575B"
			data = Data() # Instance de Data() qui va nous permettre de stocker les str récupérées grâce à main()
			self.top = tk.Toplevel()
			self.top.title("Nouvelle Commande")
			self.top.geometry("600x500")
			self.top.configure(background=color)
			self.top.bind("<Escape>", lambda d: self.top.destroy())

			
			def afficherImage(event):
				testvariable.set(event.data)
    			# get the value from string variable
				self.top.file_name=testvariable.get()
				global chemin
				chemin=self.top.file_name
				# takes path using dragged file
				image_path = Image.open(str(self.top.file_name))
				# resize image
				reside_image = image_path.resize((300, 205), Image.ANTIALIAS)
				# displays an image
				self.top.image = ImageTk.PhotoImage(reside_image)
				image_label = Label(labelframe, image=self.top.image).pack()
			
			def prendrephoto():
				global chemin
				chemin="./capture.jpg"
				cap = cv2.VideoCapture(0)
				while(True):
					ret, frame = cap.read()
					rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2BGRA)
					
					cv2.imshow('frame', rgb)
					if cv2.waitKey(1) & 0xFF == ord('q'):
						out = cv2.imwrite('capture.jpg', frame)
						break
				cap.release()
				cv2.destroyAllWindows()
			
			def enregistrer(chemin_photo):
				file=chemin_photo
				ftp = FTP('localhost')
				ftp.login(user='john',passwd='john')
				ftp.cwd("../../srv/ftp") # On se déplace vers le dossier où l'on veux stocker nos photos
				ftp.storbinary("STOR "+ os.path.basename(file), open(file, "rb"))
				print("Photo enregistrée.")
				ftp.quit()

			self.titreDepotDate = tk.Label(self.top,text="Date de depot",bg=color).grid(row=1)
			self.titreidClient = tk.Label(self.top,text="Identifiant du client",bg=color).grid(row=2)
			self.titreidCommande = tk.Label(self.top,text="Identifiant de la commande",bg=color).grid(row=3)
			self.titreAdresseLivraison = tk.Label(self.top,text="Adresse de livraison",bg=color).grid(row=4)
			self.titreCP = tk.Label(self.top,text="Code postale",bg=color).grid(row=5)
			self.titreVille = tk.Label(self.top,text="Ville",bg=color).grid(row=6)
			self.titrePrixConvenue = tk.Label(self.top,text="Prix Convenu",bg=color).grid(row=7)
			self.titreDevis= tk.Label(self.top,text="Devis",bg=color).grid(row=8)
			self.titreAcompte = tk.Label(self.top,text="Acompte",bg=color).grid(row=9)
			self.titreDescription = tk.Label(self.top, text="Description",bg=color).grid(row=10)
			self.titretravaf = tk.Label(self.top, text="Travail à faire", bg=color).grid(row=11)
			self.titreNmbcmd = tk.Label(self.top, text="Nombre de commande", bg=color).grid(row=12)


			testvariable = StringVar()
			textlabel=Label(self.top, text='Glissez la photo ici', bg='#fcba03').grid(row=13)
			entrybox = Entry(self.top, textvar=testvariable, width=20)
			entrybox.grid(row=13, column=2)
			entrybox.drop_target_register(DND_FILES)
			entrybox.dnd_bind('<<Drop>>', afficherImage)

			labelframe = LabelFrame(self.top, bg='gold')
			labelframe.grid(row=14,column=2)

			#tk.Button(self.top, text ="Valider 1", command = lambda: enregistrer(chemin)).grid(row=11)



			if type(Comm) == tuple:
				
				print(Comm)
				data.dateDepot = tk.Entry(self.top)
				data.dateDepot.insert(0, Comm[1])
				data.dateDepot.config(state='disabled')
				data.idClient = tk.Entry(self.top)
				data.idClient.insert(0, idC)
				data.idClient.config(state='disabled')
				data.idCommande = tk.Entry(self.top)
				data.idCommande.insert(0, Comm[0])
				data.adresseLivraison = tk.Entry(self.top)
				data.adresseLivraison.insert(0, Comm[2])
				data.adresseLivraison.config(state='disabled')
				data.Nmbcmd = tk.Entry(self.top)
			
				data.cp = tk.Entry(self.top)
				data.cp.insert(0, Comm[3])
				data.cp.config(state='disabled')
				data.ville = tk.Entry(self.top)
				data.ville.insert(0, Comm[4])
				data.ville.config(state='disabled')
				data.prixConvenu = tk.Entry(self.top)
				data.prixConvenu.insert(0, Comm[5])
				data.devis = tk.Entry(self.top)
				data.devis.insert(0, Comm[6])
				data.acompte = tk.Entry(self.top)
				data.acompte.insert(0, Comm[7])
				data.Description = tk.Text(self.top,height=1,width=4)
				data.Description.insert(0, Comm[9])
				data.travaf = tk.Entry(self.top)
				data.travaf.insert(0, Comm[10])


				data.dateDepot.grid(row=1, column=2)
				data.idClient.grid(row=2, column=2)
				data.idCommande.grid(row=3, column=2)
				data.adresseLivraison.grid(row=4, column=2)
				data.cp.grid(row=5, column=2)
				data.ville.grid(row=6, column=2)
				data.prixConvenu.grid(row=7, column=2)
				data.devis.grid(row=8, column=2)
				data.acompte.grid(row=9, column=2)
				data.Description.grid(row=10, column=2)
				data.travaf.grid(row=11, column=2)
				data.Nmbcmd.grid(row=12, column=2)

				testvariable = StringVar()
				textlabel=Label(self.top, text='Glissez la photo ici', bg='#fcba03').grid(row=13,)
				entrybox = Entry(self.top, textvar=testvariable, width=20)
				entrybox.grid(row=13, column=2)
				entrybox.drop_target_register(DND_FILES)
				entrybox.dnd_bind('<<Drop>>', afficherImage)

				#tk.Button(self.top, text ="Valider 2", command = lambda: enregistrer(chemin)).grid(row=14, column=2)

				labelframe = LabelFrame(self.top, bg='gold')
				labelframe.grid(row=14, column=2)				

			else:
				data.dateDepot = tk.Entry(self.top)
				data.dateDepot.insert(0, datetime.today().strftime('%d-%m-%Y'))
				data.idClient = tk.Entry(self.top)
				data.idClient.insert(0, idC)
				data.idCommande = tk.Entry(self.top)
				data.adresseLivraison = tk.Entry(self.top)
			
				if type(Comm) == list and len(Comm) > 1:
					data.adresseLivraison.insert(0, Comm[0][0])
				data.cp = tk.Entry(self.top)
				if type(Comm) == list and len(Comm) > 1:
					data.cp.insert(0, Comm[0][1])
				data.ville = tk.Entry(self.top)
				if type(Comm) == list and len(Comm) > 1:
					data.ville.insert(0, Comm[0][2])
				data.prixConvenu = tk.Entry(self.top)
				data.devis = tk.Entry(self.top)
				data.acompte = tk.Entry(self.top)
				data.Description = tk.Text(self.top,height=3,width=20)
				data.travaf = tk.Entry(self.top)
				data.Nmbcmd = tk.Entry(self.top)

				data.dateDepot.grid(row=1, column=2)
				data.idClient.grid(row=2, column=2)
				data.idCommande.grid(row=3, column=2)
				data.adresseLivraison.grid(row=4, column=2)
				data.cp.grid(row=5, column=2)
				data.ville.grid(row=6, column=2)
				data.prixConvenu.grid(row=7, column=2)
				data.devis.grid(row=8, column=2)
				data.acompte.grid(row=9, column=2)
				data.Description.grid(row=10, column=2)
				data.travaf.grid(row=11, column=2)
				data.Nmbcmd.grid(row=12, column=2)

				testvariable = StringVar()
				textlabel=Label(self.top, text='Glissez la photo ici', bg='#fcba03').grid(row=13)
				entrybox = Entry(self.top, textvar=testvariable, width=20)
				entrybox.grid(row=13, column=2)
				entrybox.drop_target_register(DND_FILES)
				entrybox.dnd_bind('<<Drop>>', afficherImage)

				#tk.Button(self.top, text ="Valider 3", command = lambda: enregistrer(chemin)).grid(row=10, column=2)

				labelframe = LabelFrame(self.top, bg='gold')
				labelframe.grid(row=14, column=2)


			# - Permet l'enregistrement de la BDD -
			def getVarCommande(data,self,idC):

				global chemin
				photos=chemin

				if data.Nmbcmd.get() == '':
					ct = 1
				else:
					ct = int(data.Nmbcmd.get())
				
				if data.idCommande.get()!="" and data.idClient.get()!="" and data.Description.get("1.0",END)!="":
					#print("idCommande: "+str(data.idCommande.get()))
					#print(str((float(data.devis.get()))-(float(data.acompte.get()))))
					conn = sqlite3.connect('base.db')
					rqt = "INSERT OR REPLACE INTO Commandes (dateDepot,idClient,idCommande,adresseLivraison,cp,ville,prixConvenu,devis,acompte,resterPaye,description,travaf,photo) VALUES ('"+data.dateDepot.get()+"','"+data.idClient.get()+"','"+data.idCommande.get()+"','"+data.adresseLivraison.get()+"','"+data.cp.get()+"','"+data.ville.get()+"','"+data.prixConvenu.get()+"','"+data.devis.get()+"','"+data.acompte.get()+"','"+str((float(data.devis.get()))-(float(data.acompte.get())))+"','"+data.Description.get("1.0",END)+"','"+data.travaf.get()+"','"+photos+"')"
					cur = conn.cursor()
					cur.execute(rqt)
					conn.commit()
					for i in range(ct):
						rqt = "INSERT OR REPLACE INTO Commandes (dateDepot,idClient,idCommande,adresseLivraison,cp,ville,prixConvenu,devis,acompte,resterPaye,description,travaf,photo) VALUES ('"+data.dateDepot.get()+"','"+data.idClient.get()+"','"+str(int(data.idCommande.get())+i)+"','"+data.adresseLivraison.get()+"','"+data.cp.get()+"','"+data.ville.get()+"','"+data.prixConvenu.get()+"','"+data.devis.get()+"','"+data.acompte.get()+"','"+str((float(data.devis.get()))-(float(data.acompte.get())))+"','"+data.Description.get("1.0",END)+"','"+data.travaf.get()+"','"+photos+"')"
						cur.execute(rqt)
						conn.commit()
					cur.close()
					conn.close()
				
				enregistrer(chemin)
				self.top.destroy()
					
			#self.top.bind("<Return>", lambda even,d=data,s=self,i=idC: getVarCommande(d,s,i))
			self.btnValide = tk.Button(self.top, text='Valider', highlightbackground=color,command=partial(getVarCommande,data,self,idC)).place(x=100,y=400)
			self.btnBon = tk.Button(self.top, text='Bon de commande', highlightbackground=color, command=partial(setCommtoBon,data))
			self.btnBon.place(x=200,y=400)
			self.btnWebcam = tk.Button(self.top, text='Prendre une photo', highlightbackground=color,command=prendrephoto).place(x=370,y=400)
			

		def fact_cli():
			from docx import Document 
			from datetime import datetime
			con = sqlite3.connect('base.db')

			cur = con.cursor()
			cur.execute('SELECT idClient FROM Commandes;')
			idCli = cur.fetchall()

			l =[]

			for i in range(len(idCli)):
				for e in range(len(idCli[i])):
					l.append(str(idCli[i][e]))
			l = list(set(l))

			comm = []
			for i in range(len(l)):
				cur.execute('SELECT * FROM Commandes WHERE idClient=?;', l[i])
				comm.append(cur.fetchall())
   
			#print(str(comm[1][0][11]) + "id: "+ str(comm[0][0][11]))

			for i in range(len(comm)):
				doc = Document()
				for e in range(len(comm[i])):
					for x in range(len(comm[i][e])):
						l.append(str(comm[i][e][x]))
					doc.add_paragraph(' '.join(l))
					l = []
				doc.save('factures/' + str(comm[i][0][11]) + '_' +datetime.today().strftime("%d-%m-%Y")+'.docx')
			showinfo(
			title='Facture',
			message='Les factures ont été générées.'
			)
        	
		def setCommtoBon(data):
			from datetime import datetime
			from docx import Document
			from os import listdir
			
			print(dir(data))
		
			print(data.idClient.get())
			
			con = sqlite3.connect('base.db')
		
		
			cur = con.cursor()
			cur.execute('SELECT * FROM Clients WHERE idClient = ?;', [data.idClient.get()])
			
			infoc = cur.fetchall()
			
			nom_entcli = infoc[0][1]
			adr_cli = data.adresseLivraison.get()
			adr_ent = infoc[0][2]
			tel_cli = infoc[0][8]
			prix = data.prixConvenu.get()
			qnt = 1

			entre = {"[date]":datetime.today().strftime('%d-%m-%Y'),"[datefdm]":"30-" + datetime.today().strftime('%m-%Y'),"[Nom de l'entreprise]":"REST'OR","[Entrecli]":nom_entcli,"[Adresse de l'entreprise du client]":adr_cli, "[SIRET du client]":"SRTcli", "[Numéro de TVA du client]":"VTA000", "[Adresse]":"ROUTE DE PONT L'EVEQUE 27260 CORMEILLES", "[SIRET]":"SRT000", "[Numéro de TVA]":"TV090909", "[Téléphone]":tel_cli, "[QNT]":qnt, "[prix]":prix}
			
			doc = Document('factures/exemple.docx')
	
			for i in entre:
				for p in doc.paragraphs:
					if p.text.find(i)>=0:
						p.text = p.text.replace(i, entre[i])
			
			
			dirs = listdir("factures/")
			dirs.sort(reverse=True)
			print(dirs)
						
			for i in range(len(dirs)):
				if "Facture" in dirs[i]:
						indice = dirs[i][len(dirs[i]) -6]
						break
				else:
						indice = 0
			doc.save('factures/Facture'+ str(int(indice) + 1)+'.docx')
			showinfo(
			title='Bon de commande',
			message='Le bon de commande à été généré.'
			)
			
			
					
		# - Permet d'ajouter un client -
		def windowClient(self):

			w = 25
			l = 12

			color = "#52575B"

			data = Data()

			self.top = tk.Toplevel(bg=color)
			self.top.title("Nouveau client")
			self.top.geometry("400x400")
			self.top.bind("<Escape>", lambda d: self.top.destroy())
			
			#self.labelTitre = tk.Label(self.top,text="Nouveaux clients",fg= color,bg=color,font=10).grid(row=0)

			self.titreidClient = tk.Label(self.top,text='Identifiant du client',bg=color).grid(row=1)
			self.titrefirmeNom = tk.Label(self.top,text='Nom Entreprise',bg=color).grid(row=2)
			self.titrefirmeAdresse = tk.Label(self.top,text='Adresse Entreprise',bg=color).grid(row=3)
			self.titrefirmeCP = tk.Label(self.top,text='Code Postale Entreprise',bg=color).grid(row=4)
			self.titrefirmeVille = tk.Label(self.top,text='Ville Entreprise',bg=color).grid(row=5)
			self.titreclientNom = tk.Label(self.top,text='Nom',bg=color).grid(row=6)
			self.titrePrenom = tk.Label(self.top,text='Prenom',bg=color).grid(row=7)
			self.titreadresseMail = tk.Label(self.top,text='Adresse Mail',bg=color).grid(row=8)
			self.titreSoldeCompte = tk.Label(self.top,text='Compte poids',bg=color).grid(row=9)
			self.titreTelephone = tk.Label(self.top,text='Tel',bg=color).grid(row=10)
			self.titreNumSiret = tk.Label(self.top,text='Numéro de siret',bg=color).grid(row=11)
			self.titreNumTVA = tk.Label(self.top,text='Numéro de TVA',bg=color).grid(row=12)


			data.idClient = tk.Entry(self.top)
			data.firmeNom = tk.Entry(self.top)
			data.firmeAdresse = tk.Entry(self.top)
			data.firmeCP = tk.Entry(self.top)
			data.firmeVille = tk.Entry(self.top)
			data.clientNom = tk.Entry(self.top)
			data.clientPrenom = tk.Entry(self.top)
			data.adresseMail = tk.Entry(self.top)
			data.soldeCompte = tk.Entry(self.top)
			data.telephone = tk.Entry(self.top)
			data.numSiret = tk.Entry(self.top)
			data.numTVA = tk.Entry(self.top)


			data.idClient.grid(row=1, column=2)
			data.firmeNom.grid(row=2, column=2)
			data.firmeAdresse.grid(row=3, column=2)
			data.firmeCP.grid(row=4, column=2)
			data.firmeVille.grid(row=5, column=2)
			data.clientNom.grid(row=6, column=2)
			data.clientPrenom.grid(row=7, column=2)
			data.adresseMail.grid(row=8, column=2)
			data.soldeCompte.grid(row=9, column=2)
			data.telephone.grid(row=10, column=2)
			data.numSiret.grid(row=11, column=2)
			data.numTVA.grid(row=12, column=2)


			# - Permet l'enregistrement de la BDD -
			def getVarClients(data,self):
				if data.idClient.get()!="" and data.firmeNom.get()!="" and data.clientNom.get()!="":
					conn = sqlite3.connect('base.db')
					rqt = "INSERT INTO Clients (idClient,firmeNom,firmeAdresse,firmeCP,firmeVille,clientNom,clientPrenom,adresseMail,telephone,soldeCompte, numSiret, numTVA) VALUES ('"+data.idClient.get()+"','"+data.firmeNom.get()+"','"+data.firmeAdresse.get()+"','"+data.firmeCP.get()+"','"+data.firmeVille.get()+"','"+data.clientNom.get()+"','"+data.clientPrenom.get()+"','"+data.adresseMail.get()+"','"+data.telephone.get()+"','"+data.soldeCompte.get()+"','"+data.numSiret.get()+"','"+data.numTVA.get()+"');"
					cur = conn.cursor()
					cur.execute(rqt)
					conn.commit()
					cur.close()
					conn.close()
					self.top.destroy()
        		
			self.top.bind("<Return>", lambda e, d=data,s=self: getVarClients(d,s))
			self.btnValide = tk.Button(self.top, text='Valider', highlightbackground =color,command=partial(getVarClients,data,self)).place(x=230,y= 330)


		def showClients(event):
			
			tree = event.widget
			citem = tree.focus()
			val = tree.item(citem)
			color1 = "#666363"
			fg = 13	
			
			print(len(val['values']))
			if len(val['values']) == 9:
				entreprise = val["values"][0]
				nom = val["values"][4]
				prenom = val["values"][5]
				compte = val["values"][3]
				tel = val["values"][7]
				mail = val["values"][6]

				self.topShowCommande = tk.Toplevel(bg=color1)
				self.topShowCommande.geometry("3000x700")

				conn = sqlite3.connect('base.db')
				cur = conn.cursor()
			
				tree = ttk.Treeview(self.topShowCommande,column=("c1", "c2", "c3","c4","c5","c6","c7","c8","c9","c10","c11"), show='headings')

			# - IL AFFICHE TOUTES LES COLONNES 
				l=["Identifiant de la commande","Date de dépôt","Adresse de livraison","Code Postale","Ville","Prix convenu", "Devis", "Acompte", "Reste à payer", "Description", "Identifiant du client"] 

				for i in range(len(l)):
					tree.column("#"+str(i+1), anchor=tk.CENTER)
					tree.heading("#"+str(i+1), text=l[i])
				tree.place(x=0,y=250)

				cur.execute('SELECT * FROM Commandes WHERE idClient = (SELECT idClient FROM Clients WHERE clientNom = ? and clientPrenom = ?);', (nom,prenom) )
				fetch = cur.fetchall()
				print(fetch)
				for res in fetch:
					tree.insert('', 'end', values=(res))
				cur.close()
				conn.close()

                # Fini ou non ??
				entreCompte = 100
				compte = str(entreCompte) + 'g'

				self.LabelFrames = tk.Label(self.topShowCommande,text=entreprise,font = fg,bg=color1).place(x=20,y=25)
				self.LabelName = tk.Label(self.topShowCommande,text=nom,bg=color1).place(x=20,y=50)
				self.LabelFirstName = tk.Label(self.topShowCommande,text=prenom,bg=color1).place(x=20,y=75)

				self.LabelMail = tk.Label(self.topShowCommande,text=mail,bg=color1).place(x=20,y=110)
				self.LabelNumTel = tk.Label(self.topShowCommande,text=tel,bg=color1).place(x=20,y=135)

				self.LabelCompte = tk.Label(self.topShowCommande,text=compte,bg=color1).place(x=340,y=135)
				self.titreCompte = tk.Label(self.topShowCommande,text='Compte poids : ',bg=color1).place(x=220,y=135)
			
		

		def windowSett():
			def destryWindowSett(self):
				self.settFrame.pack_forget()
				self.btnFlech.pack_forget()
			self.settFrame = Frame(window,bg="white",width=400,height=10000,borderwidth = 1,relief="raised",).pack(side=RIGHT)
			self.btnFlech = tk.Button(window,text="➡️",width=3,height=1,font=20,highlightbackground = "white",bg = 'grey',command=destryWindowSett(self)).place(x=1200,y=350)
			




			
		# - Affiche des boutons permettant d'accéder aux fonctionnalités -
		def prolongBtn(self):
			btnprolong5 = tk.Button(window,text="Facturation",bg="#666363",relief="ridge",highlightbackground = "#666363",command = partial
			(fact_cli),width=18)
			btnprolong5.place(x=1165,y=230)
			btnprolong2 = tk.Button(window,text="Nouveau client",bg="#666363",relief="ridge",highlightbackground = "#666363",command = partial(windowClient,self),width=18)
			btnprolong2.place(x=1165,y=110)
			btnprolong3 = tk.Button(window,text="Toutes les commandes",bg="#666363",relief="ridge",highlightbackground = "#666363",command = showAllCommandes,width=18)
			btnprolong3.place(x=1165,y=140)
			btnprolong4 = tk.Button(window,text="Tout les clients",bg="#666363",relief="ridge",highlightbackground = "#666363",command = showAllClients,width=18)
			btnprolong4.place(x=1165,y=170)	
			btnprolong1 = tk.Button(window,text="⚙️ Paramètres",bg="#666363",relief="ridge",highlightbackground = "#666363",command = windowSett,width=18)
			btnprolong1.place(x=1165,y=200)
			

		# - Bouton qui déroule les boutons permettants d'accéder aux fonctionnalités -
		btn1 = tk.Button(window,text="+",width=3,height=1,font=2,highlightbackground = "#666363",command=partial(prolongBtn,self)).place(x=1200,y=80)
		btn2 = tk.Button(window,text="Fermer",width=5,height=1,font=20,highlightbackground = "#666363",command=quit).place(x=1200,y=30)

		# - Affiche la fenêtre principale à l'écran puis attend que l'usager pose une action -
		window.mainloop()


# - Initialisation de l'UI -
app = main()